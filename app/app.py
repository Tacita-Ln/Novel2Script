"""
Flask Web 服务 - 提供小说转剧本的交互界面
"""

import io
import re
from pathlib import Path

import chardet
import docx
import yaml
from flask import Flask, render_template, request, jsonify

# 从 novel_parser 导入核心函数和全局 client
from .novel_parser import (
    parse_novel_to_script,
    check_chapters,
    client,          # 智谱客户端
    MODEL_NAME       # 模型名称
)

app = Flask(__name__)


def extract_text_from_file(file_storage):
    """
    从上传的文件中提取纯文本，支持 .txt, .md, .docx
    """
    filename = file_storage.filename
    ext = Path(filename).suffix.lower()
    content = file_storage.read()

    if ext in ('.txt', '.md'):
        detected = chardet.detect(content)
        encoding = detected.get('encoding', 'utf-8')
        if encoding.lower() in ('gb2312', 'gbk'):
            encoding = 'gbk'
        return content.decode(encoding, errors='replace')

    if ext == '.docx':
        doc = docx.Document(io.BytesIO(content))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        full_text.append(cell_text)
        return '\n'.join(full_text)

    raise ValueError(f"不支持的文件格式: {ext}，请上传 .txt, .md 或 .docx 文件")


@app.route('/')
def index():
    return render_template('index.html')


@app.route('/convert', methods=['POST'])
def convert():
    novel_text = request.form.get('novel_text', '').strip()
    uploaded_file = request.files.get('novel_file')
    force_convert = request.form.get('force_convert') == 'true'

    if uploaded_file and uploaded_file.filename != '':
        try:
            novel_text = extract_text_from_file(uploaded_file)
        except Exception as e:
            return jsonify({'error': f'文件解析失败: {str(e)}'}), 400

    if not novel_text:
        return jsonify({'error': '请输入小说内容或上传文件'}), 400

    if len(novel_text) > 80000:
        return jsonify({'error': f'小说文本过长，当前 {len(novel_text)} 字符，最大支持 80000 字符'}), 400

    if not force_convert:
        chapter_check = check_chapters(novel_text, min_chapters=3)
        if not chapter_check['valid']:
            return jsonify({
                'error': chapter_check['message'],
                'error_type': 'insufficient_chapters',
                'chapters_count': chapter_check['chapters_count']
            }), 400

    result = parse_novel_to_script(novel_text, skip_chapter_check=force_convert)
    if not result['success']:
        return jsonify({'error': result['error']}), 400

    yaml_str = yaml.dump(result['data'], allow_unicode=True, sort_keys=False)
    return jsonify({'yaml': yaml_str})


@app.route('/parse_docx', methods=['POST'])
def parse_docx():
    if 'file' not in request.files:
        return jsonify({'error': '没有文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': '文件名为空'}), 400
    try:
        text = extract_text_from_file(file)
        return jsonify({'text': text})
    except Exception as e:
        return jsonify({'error': str(e)}), 400


@app.route('/batch_convert', methods=['POST'])
def batch_convert():
    files = request.files.getlist('novel_files')
    if not files:
        return jsonify({'error': '请至少上传一个文件'}), 400

    results = []
    for file in files:
        if file.filename == '':
            continue
        try:
            novel_text = extract_text_from_file(file)
        except Exception as e:
            results.append({
                'filename': file.filename,
                'success': False,
                'error': f'文件解析失败: {str(e)}'
            })
            continue

        if len(novel_text) > 80000:
            results.append({
                'filename': file.filename,
                'success': False,
                'error': f'文本过长（{len(novel_text)}字符），最大支持80000字符'
            })
            continue

        chapter_check = check_chapters(novel_text, min_chapters=3)
        if not chapter_check['valid']:
            results.append({
                'filename': file.filename,
                'success': False,
                'error': chapter_check['message']
            })
            continue

        result = parse_novel_to_script(novel_text, skip_chapter_check=False)
        if result['success']:
            yaml_str = yaml.dump(result['data'], allow_unicode=True, sort_keys=False)
            results.append({
                'filename': file.filename,
                'success': True,
                'yaml': yaml_str
            })
        else:
            results.append({
                'filename': file.filename,
                'success': False,
                'error': result['error']
            })

    return jsonify({'results': results})


@app.route('/ask', methods=['POST'])
def ask_question():
    data = request.get_json()
    question = data.get('question', '').strip()
    current_yaml = data.get('current_yaml', '').strip()

    if not question:
        return jsonify({'error': '请输入问题'}), 400
    if not current_yaml:
        return jsonify({'error': '当前没有YAML剧本，请先转换小说'}), 400

    prompt = f"""
            你是一个专业的剧本编辑助手。用户有一个YAML格式的剧本（如下所示），现在提出了一个关于剧本的问题或修改要求。

            请遵循以下规则：
            1. 如果用户要求修改剧本（例如“增加某角色的台词”、“调整某场景的情绪”），请输出修改后的完整YAML剧本（必须严格遵循原有的YAML结构，只输出YAML代码块）。
            2. 如果用户只是提问（例如“第二场的冲突是否合理？”），请用文字回答，不要输出YAML。
            3. 如果要求修改，但改动很小，也请输出完整的修改后YAML，不要只输出diff。
            4. 输出YAML时，必须用 ```yaml ... ``` 代码块包裹。

            用户的问题：{question}

            当前剧本YAML：
            ```yaml
            {current_yaml}
            ```
            请回答：
            """

    try:
        response = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": "你是一个剧本打磨专家，根据用户问题提供修改建议或直接输出修改后的完整YAML剧本。"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            top_p=0.9,
        )
        raw_answer = response.choices[0].message.content

        # 提取 YAML 代码块（支持 ```yaml ... ``` 或 ``` ... ```）
        yaml_match = re.search(r'```yaml\n(.*?)\n```', raw_answer, re.DOTALL)
        if not yaml_match:
            yaml_match = re.search(r'```\n(.*?)\n```', raw_answer, re.DOTALL)

        if yaml_match:
            new_yaml = yaml_match.group(1).strip()
            try:
                yaml.safe_load(new_yaml)
                return jsonify({
                    'answer': raw_answer,
                    'new_yaml': new_yaml,
                    'has_yaml': True
                })
            except Exception:
                return jsonify({
                    'answer': raw_answer,
                    'new_yaml': None,
                    'has_yaml': False,
                    'warning': '检测到YAML代码块但解析失败，请手动修改'
                })
        else:
            return jsonify({
                'answer': raw_answer,
                'new_yaml': None,
                'has_yaml': False
            })
    except Exception as e:
        return jsonify({'error': f'调用模型失败: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)